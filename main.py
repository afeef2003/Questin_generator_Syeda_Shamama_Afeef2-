"""
Main implementation file for Math Question Generator
Run this file to generate questions and create Word document output
"""
import sys, os
sys.path.append(os.path.join(os.path.dirname(__file__), 'src'))

import argparse
import os
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import matplotlib.patches as patches
from question_generator import MathQuestionGenerator



def generate_geometry_image(rows, cols, radius, filename):
    """Generate visual representation for geometry questions"""
    fig, ax = plt.subplots(1, 1, figsize=(8, 6))
    for i in range(rows):
        for j in range(cols):
            x = j * (2 * radius) + radius
            y = (rows - 1 - i) * (2 * radius) + radius
            circle = patches.Circle((x, y), radius,
                                    facecolor='lightblue',
                                    edgecolor='navy',
                                    linewidth=2,
                                    alpha=0.8)
            ax.add_patch(circle)
    rect = patches.Rectangle((0, 0), cols * 2 * radius, rows * 2 * radius,
                              linewidth=3, edgecolor='red', facecolor='none',
                              linestyle='--', alpha=0.7)
    ax.add_patch(rect)
    ax.set_xlim(-0.5, cols * 2 * radius + 0.5)
    ax.set_ylim(-0.5, rows * 2 * radius + 0.5)
    ax.set_aspect('equal')
    ax.set_title('Top View of Tightly Packed Objects', fontsize=14, fontweight='bold')
    ax.grid(True, alpha=0.3)
    ax.set_xticks([])
    ax.set_yticks([])
    plt.tight_layout()
    plt.savefig(filename, dpi=300, bbox_inches='tight', facecolor='white')
    plt.close()
    print(f"Generated image: {filename}")

def create_word_document(questions, output_file):
    """Create Word document with generated questions"""
    doc = Document()
    title = doc.add_heading('Generated Math Assessment Questions', 0)
    title.alignment = 1
    doc.add_paragraph()
    desc_para = doc.add_paragraph()
    desc_para.add_run("Assessment Description: ").bold = True
    desc_para.add_run("This assessment contains multiple-choice questions designed to test students' problem-solving abilities in counting arrangements and geometric spatial reasoning.")
    doc.add_page_break()

    for i, question in enumerate(questions, 1):
        doc.add_heading(f'Question {i}', level=1)
        q_para = doc.add_paragraph()
        q_para.add_run("Question: ").bold = True
        q_para.add_run(question['question'])

        if question.get('image_path') and os.path.exists(question['image_path']):
            doc.add_paragraph()
            doc.add_picture(question['image_path'], width=Inches(5))
            doc.add_paragraph()

        doc.add_paragraph().add_run("Options:").bold = True
        option_letters = ['A', 'B', 'C', 'D', 'E']
        for j, option in enumerate(question['options']):
            opt_para = doc.add_paragraph(f"({option_letters[j]}) {option}")
            if j == question['correct_index']:
                for run in opt_para.runs:
                    run.bold = True

        doc.add_paragraph()
        exp_para = doc.add_paragraph()
        exp_para.add_run("Explanation: ").bold = True
        exp_para.add_run(question['explanation'])

        doc.add_paragraph()
        curr_para = doc.add_paragraph()
        curr_para.add_run("Curriculum Mapping:").bold = True
        doc.add_paragraph(f"Subject: {question['subject']}")
        doc.add_paragraph(f"Unit: {question['unit']}")
        doc.add_paragraph(f"Topic: {question['topic']}")

        if i < len(questions):
            doc.add_page_break()

    doc.save(output_file)
    print(f"Word document saved: {output_file}")

def generate_formatted_output(questions, output_file):
    """Generate questions in the specified format"""
    with open(output_file, 'w') as f:
        f.write("@title Middle School Mathematics Problem Solving Assessment\n")
        f.write("@description This assessment contains multiple-choice questions designed to test students' problem-solving abilities in counting arrangements and geometric spatial reasoning.\n\n")

        for i, question in enumerate(questions, 1):
            f.write(f"@question {question['question']}\n")
            f.write("@instruction Choose the correct answer.\n")
            f.write("@difficulty moderate\n")
            f.write(f"@Order {i}\n")
            for j, option in enumerate(question['options']):
                if j == question['correct_index']:
                    f.write(f"@@option {option}\n")
                else:
                    f.write(f"@option {option}\n")
            f.write("@explanation\n")
            f.write(f"{question['explanation']}\n")
            f.write(f"@subject {question['subject']}\n")
            f.write(f"@unit {question['unit']}\n")
            f.write(f"@topic {question['topic']}\n")
            f.write("@plusmarks 1\n\n")

    print(f"Formatted output saved: {output_file}")

def main():
    parser = argparse.ArgumentParser(description='Generate math questions')
    parser.add_argument('--output', '-o', default='generated_questions.docx',
                        help='Output Word document filename')
    parser.add_argument('--count', '-c', type=int, default=2,
                        help='Number of questions to generate')
    parser.add_argument('--images-dir', default='images',
                        help='Directory to save generated images')
    parser.add_argument('--mode', choices=['dynamic', 'fixed'], default='dynamic',
                        help='Choose question mode: dynamic (generator) or fixed (predefined)')
    args = parser.parse_args()

    os.makedirs(args.images_dir, exist_ok=True)
    os.makedirs('output', exist_ok=True)

    questions = []
    if args.mode == 'fixed':
        print("Loading fixed questions...")
        questions = get_fixed_questions()
        # Truncate or pad to match requested count
        questions = questions[:args.count]
    else:
        print("Generating dynamic questions...")
        generator = MathQuestionGenerator()
        for i in range(1, args.count + 1):
            if i % 2 == 1:
                questions.append(generator.generate_counting_question())
            else:
                geo_q = generator.generate_geometry_question()
                img_file = os.path.join(args.images_dir, f'geometry_question_{i//2}.png')
                generate_geometry_image(2, 4, 1.5, img_file)
                geo_q['image_path'] = img_file
                questions.append(geo_q)

    output_docx = os.path.join('output', args.output)
    output_txt = os.path.join('output', 'formatted_questions.txt')

    create_word_document(questions, output_docx)
    generate_formatted_output(questions, output_txt)

    print("\n✅ Generation complete!")
    print(f"📄 Word document: {output_docx}")
    print(f"📝 Formatted text: {output_txt}")
    print(f"🖼️  Images directory: {args.images_dir}")

if __name__ == "__main__":
    main()
