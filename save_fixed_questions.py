import os
from docx import Document
from src.question_generator import MathQuestionGenerator

def save_fixed_questions_to_word(output_file="fixed_questions.docx"):
    gen = MathQuestionGenerator()
    questions = gen._load_fixed_questions()  # get the fixed questions

    doc = Document()
    doc.add_heading("Fixed Math Questions", 0).alignment = 1  # Centered title

    for i, q in enumerate(questions, 1):
        doc.add_heading(f"Question {i}", level=1)

        # Question text
        q_para = doc.add_paragraph()
        q_para.add_run("Question: ").bold = True
        q_para.add_run(q['question'])

        # Options
        options_para = doc.add_paragraph()
        options_para.add_run("Options:").bold = True
        option_letters = ['A', 'B', 'C', 'D', 'E']
        for j, option in enumerate(q['options']):
            opt_para = doc.add_paragraph(f"({option_letters[j]}) {option}")
            if j == q['correct_index']:
                for run in opt_para.runs:
                    run.bold = True

        # Explanation
        exp_para = doc.add_paragraph()
        exp_para.add_run("Explanation: ").bold = True
        exp_para.add_run(q['explanation'])

        # Curriculum
        curr_para = doc.add_paragraph()
        curr_para.add_run("Curriculum Mapping:").bold = True
        doc.add_paragraph(f"Subject: {q['subject']}")
        doc.add_paragraph(f"Unit: {q['unit']}")
        doc.add_paragraph(f"Topic: {q['topic']}")

        if i < len(questions):
            doc.add_page_break()

    doc.save(output_file)
    print(f"✅ Fixed questions saved to Word: {output_file}")

if __name__ == "__main__":
    os.makedirs("output", exist_ok=True)
    save_fixed_questions_to_word(os.path.join("output", "fixed_questions.docx"))
